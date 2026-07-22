#!/usr/bin/env python3
"""Build paper.docx (Word manuscript) with figures and tables inline."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

FIG_DIR = "figures"


def add_caption(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + " ")
    r.bold = True
    p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.size = Pt(9)


def add_figure(doc, path, caption_label, caption_text, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    add_caption(doc, caption_label, caption_text)


def add_table(doc, caption_label, caption_text, header, rows):
    add_caption(doc, caption_label, caption_text)
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        c.text = h
        for par in c.paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = str(val)
            for par in c.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


doc = Document()

# ---- Title ----
t = doc.add_heading(
    "mapgis2shp: an open-source Python reader for the MapGIS 6.x/67 binary "
    "vector formats, validated against the native export", level=0)

# ---- Authors / affiliation ----
p = doc.add_paragraph()
p.add_run("Shijie Li*, Haiyang He*, Xu Sun, Haoyang Qin, Xiaoyu Liu, "
          "Yilin Feng, Zengyun Zuo").bold = True
doc.add_paragraph("Xi'an Mineral Resources Survey, China Geological Survey, "
                  "Xi'an, China")
doc.add_paragraph("* Co-corresponding authors. Emails: 1045105061@qq.com (S. Li); "
                  "hehaiyangcugb@foxmail.com (H. He)")

# ---- Abstract ----
doc.add_heading("Abstract", level=1)
body(doc,
"MapGIS is a closed-source geographic information system that dominates "
"Chinese geological surveying, yet its native binary vector formats — .wt "
"(point), .wl (line), and .wp (polygon) — remain undocumented and cannot be "
"read by any mainstream open-source geospatial tool, including GDAL/OGR and "
"QGIS. Survey data held in these formats are therefore effectively locked "
"out of open-source analysis pipelines unless the proprietary software is "
"used. We present mapgis2shp (PyPI distribution name mapgis2shp, import name "
"pymapgis): an open-source, Apache-2.0 Python reader that reverse-engineers "
"the MapGIS 6.x/67 binary layout, infers the coordinate reference system "
"(CRS) from the file's projection and ellipsoid index codes, reconstructs "
"polygon topology from the arc–node structure, and returns standard "
"geopandas.GeoDataFrame objects or ESRI shapefiles. Several open-source "
"MapGIS readers already exist, but none has been independently validated "
"against the reference implementation, none infers the CRS, and none has "
"been peer-reviewed. We address all three gaps through a reproducible "
"cross-validation against the official MapGIS software, on two tiers. "
"First, across 36 real geological-survey layers comprising 16,874 features, "
"feature counts, attribute schemas, and geometries are identical (point "
"distance, line Hausdorff distance, and polygon intersection-over-union all "
"indicate exact coincidence within 1e-7 degrees), and attribute values are "
"semantically equivalent in 99.9995% of 95,006 compared cells. Second, on a "
"400 MB polygon file of 78,873 features, the bounding box, coordinate "
"reference system, and the 16 native attribute fields match the official "
"export exactly, the reader produces zero invalid geometries against three "
"in the reference, and spatial coverage agrees with an intersection-over-"
"union of 99.73%. The only deviations are attributable to lossy rounding in "
"the native shapefile export and to a quantified 1.66% overlap artefact of "
"the heuristic arc-merge reconstruction at shared boundaries — an effect "
"that shrinks to a 0.27% real coverage difference once overlaps are unioned "
"away. The reader in fact preserves source data more faithfully than the "
"native export, which truncates floating-point fields and strips stored "
"string padding. mapgis2shp is distributed on PyPI and archived under a "
"citable DOI.")
doc.add_paragraph(
    "Keywords: MapGIS; reverse engineering; vector format; GeoPandas; "
    "shapefile; geological survey; open source").runs[0].italic = True

# ---- 1 Introduction ----
h1(doc, "1. Introduction")
h2(doc, "1.1 Motivation")
body(doc,
"MapGIS (Zondy Cyber, Wuhan) is the de facto geographic information system "
"of the Chinese geological-survey community, underpinning regional geological "
"mapping, mineral exploration, and engineering geology [15,16]. Its native "
"vector storage — the 6.x/67 generation's .wt (point), .wl (line), and .wp "
"(polygon) binary files — is closed: the vendor does not publish the on-disk "
"layout, and the format is distinct from the later MapGIS K9/10 \"open data "
"format\". Consequently, survey data held in these files cannot be read by "
"the open-source geospatial stack (GDAL/OGR, GeoPandas, QGIS, GRASS) without "
"the proprietary software. This lock-in impedes reproducible analysis, data "
"sharing, and long-term archival — issues repeatedly raised for proprietary "
"geoscience formats in the open-science literature.")
h2(doc, "1.2 The interoperability gap")
body(doc,
"We verified that GDAL/OGR ships no MapGIS driver: the official vector-"
"driver index contains no mapgis entry, and the OSGeo/gdal issue tracker and "
"pull requests contain no mention of \"MapGIS\" or the vendor name "
"\"ZondyCyber\" across the entire organisation [2]. Because QGIS, GRASS, "
"SAGA, and WhiteboxTools all rely on GDAL/OGR for vector I/O [1], none of "
"them can open .wt/.wl/.wp files natively. The only exchange routes "
"documented by the vendor are an ASCII \"plain-code\" format and the newer "
"MapGIS 10 open format — neither of which is the legacy binary format that "
"holds the bulk of existing survey data.")
h2(doc, "1.3 Related work")
body(doc,
"Several open-source MapGIS readers predate this work and warrant explicit "
"acknowledgement. The present software, mapgis2shp, is the matured and "
"validated successor to the first author's earlier single-file reader "
"pymapgis [3] (2022), which first published a byte-level format description; "
"the present paper adds coordinate-reference inference, a vectorised parser, "
"polygon-topology reconstruction, and — critically — independent validation "
"against the reference implementation. MathsionYang/MapGIS2ArcGIS [4] (2019, "
"C#) independently parses the .wl structure at the record level. "
"BenChao1998/ConvertMapGIS [5] (2025) is an independent third-party "
"derivative of pymapgis — not authored by the present authors — that "
"republishes a detailed byte-level specification in its README and adds a "
"graphical converter. WenboWong/OGC.net [6] (2021, C#) and earlier C# "
"format-bridge tools are further examples. Chinese-language academic "
"literature has also analysed the format programmatically — notably Wang "
"[7] on reading the .wt point file, Wu and Feng [8] on MapGIS–Geodatabase "
"conversion, and Chen [9] on MapGIS format conversion in MATLAB — but these "
"are short application notes: none ships an open-source reader, none "
"publishes a byte-level specification, and none validates output correctness.")
body(doc,
"Methodologically, this work belongs to a well-established genre of "
"open-source readers for closed or under-documented scientific binary "
"formats, including sas7bdat for SAS files [10], python-ags4 for AGS4 "
"geotechnical files [11], and the long-running GDAL library [1] for "
"numerous proprietary GIS formats. GeoSciML provides a comparable "
"geoscience-format specification effort [13], and van den Bos [14] "
"discusses binary-format reverse-engineering methodology in general terms. "
"Within the target journal, format-bridge software such as ArcGMT [12] "
"establishes precedent for this type of contribution. Crucially, no prior "
"MapGIS reader offers (i) independent validation against the reference "
"implementation, (ii) coordinate-reference inference, or (iii) a "
"reproducible validation harness. The present work closes these three gaps.")
h2(doc, "1.4 Contributions")
body(doc,
"This paper makes the following verified contributions: (1) the first "
"peer-reviewed software paper on a MapGIS 6.x/67 reader (Crossref, arXiv, "
"OpenAlex, and a manual CNKI/Wanfang search return no such paper); (2) the "
"first reader independently validated against the reference implementation "
"(the official MapGIS export) with quantified fidelity, on two tiers — 36 "
"small and medium layers (16,874 features; exact geometric coincidence; "
"99.9995% attribute equivalence) and one 400 MB file (78,873 features; "
"99.73% coverage IoU); (3) the first reader that infers and attaches the "
"CRS — the native shapefile export omits the .prj, whereas mapgis2shp "
"reconstructs a PROJ string (Krassovsky 1940 with towgs84 for the test "
"data); and (4) a reproducible, open-source (Apache-2.0) validation harness "
"distributed on PyPI, with a 36-file regression baseline. We do not claim "
"novelty for the reader, the byte-level specification, or shapefile "
"conversion as such; these are prior art.")

# ---- 2 Background + Figure 1 ----
h1(doc, "2. Background: the MapGIS 6.x/67 binary formats")
body(doc,
"All three file types share a common top-level organisation: an 8-byte "
"magic header (WMAP·D22 for points, D21 for lines, D23 for polygons), a "
"4-byte file-identifier int32, and a 4-byte data_start offset that locates "
"a 10-entry index table. Each index entry occupies 10 bytes, of which the "
"first 8 hold a (start, volume) pair describing one sub-data region. "
"Numeric values are little-endian, and strings are GBK-encoded. The point "
"section stores 93-byte records, with X and Y as little-endian doubles at "
"offsets 7–14 and 15–22 respectively. The line and polygon files share a "
"57-byte arc-index record (point count at offset 10–13; byte offset into "
"the coordinate section at 14–17) followed by a coordinate section of "
"16-byte XY pairs. The polygon file additionally stores a 24-byte topology "
"record whose bytes 8–15 carry the left and right polygon identifiers of "
"each arc, from which closed rings, shells, and holes are reconstructed. "
"CRS information resides at fixed offsets in the file header: projection "
"type at byte 109, ellipsoid at byte 110, scale denominator (double) at "
"byte 143, and — for projected CRSes — the central meridian and standard "
"parallels, encoded as DDDMMSS.sss doubles, from byte 151 onward. The "
"complete byte-level specification accompanies the software as a reference "
"document.")
add_figure(doc, FIG_DIR + "/figure1_format_layout.png",
    "Figure 1.",
    "Byte-level layout of the MapGIS 6.x/67 vector formats. (a) Common file "
    "header: 8-byte magic (WMAP·D22/D21/D23 for point/line/polygon), file "
    "identifier, data_start offset, the coordinate-reference-system bytes "
    "(projection@109, ellipsoid@110, scale@143, central meridian@151 in "
    "DDDMMSS.sss), and the 10-entry index area located at data_start. (b) "
    "Type-specific record sections: .wt point coordinate records (93 B; X/Y "
    "doubles at offsets 7–14/15–22), .wl line index (57 B) plus 16-B XY "
    "coordinate pairs, and .wp arc index (57 B), coordinate section, and "
    "24-B topology records (left/right polygon identifiers at 8–11/12–15). "
    "Strips are schematic and not to scale; each block reports its record "
    "size. All integers are little-endian; strings are GBK.")

# ---- 3 Design + Figure 2 ----
h1(doc, "3. Design and Implementation")
h2(doc, "3.1 Architecture")
body(doc,
"mapgis2shp is organised as a thin pipeline: binary I/O → record model → "
"Shapely geometries → GeoPandas GeoDataFrame. A Reader class parses the "
"file in full on construction and exposes .geodataframe, .fields, .crs, "
".bbox, and a .to_file() passthrough. A small command-line interface "
"(pymapgis input.wp output.shp) covers the common conversion case. The "
"package targets Python ≥ 3.9 and depends only on geopandas, numpy, pandas, "
"pyproj, and shapely.")
h2(doc, "3.2 Vectorised binary parsing")
body(doc,
"Binary records are decoded using NumPy structured dtypes rather than "
"per-field struct.unpack calls. The point, arc-index, and topology tables "
"are each materialised in a single np.frombuffer call, and coordinate "
"arrays are reshaped to (n,2) and scaled in one vectorised multiply. As a "
"result, parsing the largest test polygon (608 features, 2,166 arcs) takes "
"~0.33 s on a laptop. Attribute-table field descriptors are unpacked in "
"bulk through a precompiled struct.Struct that mirrors the 39-byte "
"descriptor layout.")
h2(doc, "3.3 Coordinate-reference inference")
body(doc,
"From the projection index (byte 109) and the ellipsoid index (byte 110), "
"mapgis2shp constructs a PROJ string. Longitude/latitude files (projection "
"0) yield a geographic CRS; Transverse Mercator (5), Albers (2), and "
"Lambert (3) yield projected CRSes, with the central meridian decoded "
"arithmetically from the DDDMMSS.sss double (preserving fractional seconds "
"and handling western longitudes correctly). Supported ellipsoid codes "
"include Krassovsky 1940 (Beijing 1954), Xi'an 1980, WGS84, WGS72, and "
"CGCS2000. When the ellipsoid code is unrecognised or the scale is zero, "
"the reader returns an empty CRS rather than guessing, so that downstream "
"code can detect the gap.")
h2(doc, "3.4 Polygon topology reconstruction")
body(doc,
"Polygon geometries are rebuilt from the arc–node topology. Arcs whose "
"left or right polygon identifier equals the target identifier are oriented "
"consistently, chained into closed rings by endpoint matching, and "
"classified into shells and holes by containment. Ring assembly uses a "
"spatial hash on arc endpoints (cell size 1e-5 map units) with a two-sided "
"greedy walk and a closure-competition rule, reducing the merge from "
"O(n^3) (a fresh all-pairs distance matrix at every iteration) to O(n) in "
"the common case while preserving the semantics of the original nearest-"
"pair heuristic. Rings that fail to close cleanly — a known artefact of "
"topological gaps in the source data — are passed through shapely.make_valid "
"so that the output satisfies the OGC simple-features specification; "
"degenerate rings with fewer than four coordinates are skipped.")
add_figure(doc, FIG_DIR + "/figure2_architecture.png",
    "Figure 2.",
    "mapgis2shp reader architecture and data flow. (a) Main pipeline (solid "
    "arrows): MapGIS .wt/.wl/.wp closed binary → binary I/O and record model "
    "(NumPy structured dtypes) → Shapely geometries → open outputs "
    "(GeoDataFrame / shapefile / GeoJSON). Dashed bypasses: CRS inference "
    "(projection/ellipsoid index codes → PROJ string), feeding both the "
    "parser and the output projection; and polygon topology reconstruction "
    "(arc–node → rings → shells/holes + make_valid), feeding polygon "
    "geometry. The Reader API and CLI form a thin access layer over the "
    "whole pipeline. (b) Reproducible verification: the cross-validation "
    "harness (36-layer 1:1 comparison against the official MapGIS export — "
    "geometry 100%, attributes 99.9995%; 400 MB coverage-equivalence test — "
    "IoU 99.73%) and the regression baseline (pymapgis_baseline.json + "
    "pytest), rising as dashed \"verify\" arrows into the pipeline.")

# ---- 4 Validation + Tables 1,2,3,5 ----
h1(doc, "4. Validation")
body(doc,
"The validation section is the core of this contribution: it quantifies "
"how closely the output of mapgis2shp matches the reference implementation "
"(official MapGIS) on real production data.")
h2(doc, "4.1 Dataset")
body(doc,
"The validation set comprises 36 MapGIS vector layers drawn from 1:50,000 "
"geological-survey sheets in the Kurgan region (J43C001002, Xinjiang, "
"China): 5 point layers, 18 line layers, and 13 polygon layers, totalling "
"16,874 features (Table 1). The layers span geological boundaries, faults, "
"attitudes, hydrology, and Quaternary geology, and range from small "
"(2-feature) to large (6,981-feature) files.")
add_table(doc, "Table 1.", "Validation dataset (36 layers; abbreviated).",
    ["Layer", "Type", "Features"],
    [["LDLYAAI002", "point", "242"],
     ["LDLYAAE001", "line", "1,143"],
     ["LDZOFBB001", "polygon", "608"],
     ["LFZYBCT001", "point", "6,981"],
     ["LDZOFBA002", "line", "2,222"],
     ["(31 further layers)", "—", "—"],
     ["Total", "5 pt / 18 ln / 13 poly", "16,874"]])
h2(doc, "4.2 Protocol")
body(doc,
"Each layer was processed in two ways: (i) read with mapgis2shp to produce "
"a GeoDataFrame, and (ii) exported to ESRI shapefile using the official "
"MapGIS 6.7 software (Zondy Cyber) under default export settings. The two "
"outputs were aligned by the unique ID attribute and compared along four "
"axes: Count (feature counts); Schema (attribute column-name sets); "
"Geometry (point-to-point Euclidean distance; line Hausdorff distance; "
"polygon intersection-over-union, with coincidence at tolerance 1e-7 "
"degrees); and Attributes (per-field semantic equivalence, with a "
"classifier distinguishing exact matches from native-export lossy cases "
"and genuine mismatches). The full protocol is implemented in "
"cross_validate_native.py, shipped with the software and fully reproducible.")
h2(doc, "4.3 Results")
body(doc,
"Across all 36 layers and 16,874 features: feature counts are identical "
"for 36/36 layers (16,874 = 16,874); schemas are identical for 36/36 "
"layers; geometry matches for 100% of aligned features across all 36 "
"layers (point distances, line Hausdorff distances, and polygon 1−IoU all "
"zero within 1e-7 degrees — geometries are exactly coincident); attributes "
"are semantically equivalent in 99.9995% of the 95,006 compared cells, "
"with only 2 cells differing, both floating-point fields smaller than 1e-6 "
"that the native export rounds coarsely (e.g., 6.69e-7 in mapgis2shp "
"versus 1e-6 in the native export) — no genuine semantic attribute "
"difference was found; and CRS: mapgis2shp reconstructs and attaches the "
"CRS (Krassovsky 1940 with towgs84=15.8,-154.4,-82.3,…), whereas the "
"native shapefile export produces no .prj file, so the CRS is lost.")
add_table(doc, "Table 2.", "Cross-validation results (summary across 36 layers).",
    ["Metric", "Result"],
    [["Layers compared", "36"],
     ["Features (mapgis2shp / native)", "16,874 / 16,874"],
     ["Count match", "36 / 36"],
     ["Schema match", "36 / 36"],
     ["Geometric coincidence", "100% (36 / 36 layers)"],
     ["Attribute equivalence", "99.9995% (2 native-rounding deviations)"],
     ["CRS attached", "mapgis2shp yes; native no"]])
h2(doc, "4.4 The native export is lossy; mapgis2shp is not")
body(doc,
"In 10,245 of the 95,006 compared cells (10.8%), mapgis2shp preserves "
"source information that the native shapefile export discards (Table 3): "
"(1) floating-point truncation — the native export rounds floats to 2–6 "
"significant figures (e.g., 0.6218423 → 0.62; 4.88e-6 → 5e-6), whereas "
"mapgis2shp preserves full double precision; (2) numeric formatting — "
"values stored as floating-point are exported by MapGIS as integer strings "
"in some fields (e.g., 72.0 versus 72), whereas mapgis2shp preserves the "
"stored type; (3) string padding — the native export strips leading and "
"trailing whitespace that the binary field actually stores (e.g., "
"'113000   ' → '113000'), whereas mapgis2shp preserves the raw bytes. "
"mapgis2shp is therefore not merely equivalent to the native export; it is "
"a strictly more faithful representation of the source binary data.")
add_table(doc, "Table 3.", "Native-export lossy artefacts.",
    ["Category", "Example (mapgis2shp → native)", "Cells"],
    [["Float truncation", "0.6218423 → 0.62", "~9,000"],
     ["Sci-notation rounding", "4.88e-6 → 5e-6", "~1,200"],
     ["Numeric formatting", "72.0 → 72", "310"],
     ["Whitespace stripping", "'113000   ' → '113000'", "~200"],
     ["Total native-lossy cells", "", "10,245"]])
h2(doc, "4.5 Extreme-scale validation")
body(doc,
"To stress the reader beyond the 36 survey layers, a single 400 MB MapGIS "
"polygon file (78,873 features) was compared against an official MapGIS "
"export of the same source. Because the official export had been "
"post-processed (adjacent polygons dissolved and re-coloured), a 1:1 "
"feature-wise comparison is impossible; coverage equivalence was assessed "
"through geometric unions, with areas computed on the Krassovsky ellipsoid "
"using pyproj.Geod. Results: bounding box exact match (difference = 0); "
"CRS correctly inferred as Beijing-1954 / Krassovsky (longlat) — the "
"official export had lost its .prj file; the 16 native attribute fields "
"identical (the official export carries two extra fields, strat_code and "
"ColorCode, artefacts of its post-processing); invalid geometries: "
"mapgis2shp 0, official 3; coverage equivalence: union symmetric "
"difference 5,649.94 km² against 2.11e6 km² of coverage, giving a coverage "
"IoU of 99.73%. The summed polygon areas differ by 1.66% (mapgis2shp "
"larger), an overlap artefact of the greedy arc-merge reconstruction at "
"shared boundaries (sliver bridging across topological gaps), not a "
"coverage error: once overlaps are removed by the union, the real coverage "
"difference is only 0.27%. Table 5 summarises the extreme-scale comparison.")
add_table(doc, "Table 5.",
    "Extreme-scale (400 MB, 78,873 features) validation against the official MapGIS export.",
    ["Metric", "mapgis2shp / Official", "Agreement"],
    [["File size", "400 MB / 327 MB shp", "—"],
     ["Features", "78,873 / 67,039 (dissolved)", "counts differ (post-processing)"],
     ["Bounding box", "[73.486, 32.0, 111.000, 48.0]", "exact (diff = 0)"],
     ["CRS", "Krassovsky / Beijing-54 (inferred)", "correct"],
     ["Native attribute fields", "16 / 16 (+2 post-process)", "16 identical"],
     ["Invalid geometries", "0 / 3", "reader cleaner"],
     ["Union symmetric difference", "— / 5,649.94 km²", "0.27% of coverage"],
     ["Coverage IoU", "—", "99.73%"],
     ["Summed-area gap", "+1.66%", "overlap artefact"]])

# ---- 5 Comparison + Table 4 ----
h1(doc, "5. Comparison with prior tools")
body(doc,
"Table 4 compares mapgis2shp with the prior open-source MapGIS readers and "
"the Chinese academic precedents along the axes that distinguish this work: "
"open-source licence, language, byte-level specification, CRS inference, "
"independent validation against the reference implementation, PyPI "
"distribution, and peer-reviewed publication. mapgis2shp is the only entry "
"that infers the CRS, the only one validated against the official export, "
"and the only one distributed as an installable package with a reproducible "
"validation harness.")
add_table(doc, "Table 4.", "Feature comparison with prior tools.",
    ["Tool", "Year", "Lang", "Licence", "Spec", "CRS", "Validated", "Peer-rev."],
    [["pymapgis (own)", "2022", "Python", "none", "yes", "no", "no", "no"],
     ["MathsionYang", "2019", "C#", "—", "no", "no", "no", "no"],
     ["ConvertMapGIS", "2025", "Python", "GPL-3.0", "yes", "no", "no", "no"],
     ["OGC.net", "2021", "C#", "none", "no", "no", "no", "no"],
     ["Wang (2013)", "2013", "—", "n/a", "no", "no", "no", "note"],
     ["mapgis2shp", "2026", "Python", "Apache-2.0", "yes", "yes", "yes", "this"]])

# ---- 6 Discussion ----
h1(doc, "6. Discussion")
h2(doc, "6.1 Limitations")
body(doc,
"The reader supports the most common ellipsoid codes (Krassovsky, Xi'an "
"1980, WGS84, WGS72, CGCS2000) and the most common projections "
"(geographic, Transverse Mercator, Albers, Lambert); unrecognised codes "
"yield an empty CRS that downstream code must handle. The legacy MapGIS "
"K9/10 formats are not supported, as they use a different binary layout. "
"CRS inference relies on the file's index codes rather than a full "
"parameter block, so unusual projections may be under-described. Polygon "
"topology reconstruction is heuristic: although make_valid guarantees "
"valid output, rare topological gaps in the source data can produce "
"repaired rather than exact polygons. At extreme scale (the 400 MB test), "
"the greedy arc-merge introduces a quantified 1.66% overlap artefact in "
"summed polygon areas at shared boundaries — a coverage difference of only "
"0.27% once overlaps are unioned away, but a known limitation when an "
"exact planar partition is required. A planar-enforcement post-process "
"(e.g., shapely.union of all polygons) would eliminate the overlaps at the "
"cost of runtime. Finally, the reader is read-only.")
h2(doc, "6.2 Generalisability")
body(doc,
"The validation methodology — reverse-engineering a proprietary reader and "
"then cross-validating it against the reference implementation on real "
"data, with quantified geometric and attribute fidelity — transfers "
"directly to other closed geoscience formats. The cross-validation script "
"is format-agnostic in structure and could anchor similar studies for any "
"format for which a reference implementation exists.")
h2(doc, "6.3 Open-science implication")
body(doc,
"By making MapGIS 6.x/67 data readable from the open-source Python stack, "
"mapgis2shp unlocks a large body of Chinese geological-survey data for "
"reproducible research, for integration with GDAL/GeoPandas/QGIS workflows, "
"and for long-term archival. The CRS-inference capability is of particular "
"practical importance: users who currently rely on the native export "
"silently lose the coordinate reference.")
h2(doc, "6.4 Literature coverage")
body(doc,
"The Crossref, OpenAlex, arXiv, and DOAJ indices contain no peer-reviewed "
"MapGIS reader paper. A manual search of Chinese-language databases (CNKI, "
"Wanfang) was also carried out: the retrieved literature consists almost "
"entirely of MapGIS–ArcGIS conversion-method studies and application notes "
"built on the vendor SDK or the documented ASCII plain-code format, with "
"effectively no open-source conversion tool and no byte-level specification "
"accompanied by validation. The open-source prior art enumerated in "
"Section 1.3 is fully covered, and no further prior art affecting the "
"novelty claims was identified.")

# ---- 7 Conclusions ----
h1(doc, "7. Conclusions")
body(doc,
"mapgis2shp is an open-source Python reader for the closed MapGIS 6.x/67 "
"binary vector formats. Across 36 real geological-survey layers (16,874 "
"features), its output is geometrically identical to and attribute-"
"equivalent to the official MapGIS export (99.9995% semantic equivalence, "
"with zero genuine mismatches), and on a 400 MB file (78,873 features) it "
"attains 99.73% coverage equivalence while producing fewer invalid "
"geometries than the reference. It preserves source data more faithfully "
"than the native export, which truncates floating-point fields and strips "
"stored string padding, and it additionally reconstructs the coordinate "
"reference system that the native export discards. The software and its "
"reproducible validation harness are distributed on PyPI under the "
"Apache-2.0 licence.")

# ---- Availability ----
h1(doc, "Availability and Requirements")
for lab, val in [
    ("Software name:", "mapgis2shp (import name pymapgis)"),
    ("Version:", "2.0.6"),
    ("PyPI:", "pip install mapgis2shp"),
    ("Source code:", "https://github.com/leecugb/mapgis2shp"),
    ("Archived DOI:", "https://doi.org/10.5281/zenodo.21487339"),
    ("License:", "Apache-2.0"),
    ("Operating systems:", "Windows, Linux, macOS"),
    ("Dependencies:", "Python ≥ 3.9; geopandas, numpy, pandas, pyproj, shapely"),
]:
    p = doc.add_paragraph()
    p.add_run(lab + " ").bold = True
    p.add_run(val)

h1(doc, "Data Availability Statement")
body(doc,
"The 36 MapGIS validation layers and the 400 MB extreme-scale file are "
"real geological-survey data from the China Geological Survey and are not "
"redistributed with the package. The reproducible cross-validation scripts "
"(cross_validate_native.py for the 36-layer 1:1 comparison; "
"cross_validate.py for the large-file coverage-equivalence test), the "
"per-layer results (cross_validation_report.csv), the large-file results "
"(cross_validate_large_report.md), the regression baseline "
"(pymapgis_baseline.json), and synthetic minimal test fixtures (tests/) "
"are available in the mapgis2shp repository at "
"https://github.com/leecugb/mapgis2shp and in the archived release at "
"https://doi.org/10.5281/zenodo.21487339.")

h1(doc, "CRediT author statement")
body(doc,
"S.L. conceived and designed the study, developed the software, carried "
"out the validation, and wrote the main manuscript text; S.L. also "
"prepared Figures 1 and 2. H.H. contributed to the validation and reviewed "
"the manuscript. X.S., H.Q., X.L., Y.F. and Z.Z. contributed to data "
"investigation and reviewed the manuscript. All authors reviewed and "
"approved the final manuscript.")

h1(doc, "Declaration of competing interests")
body(doc, "The authors declare no competing interests.")

h1(doc, "Funding")
body(doc,
"This work was supported by the Deep Earth Probe and Mineral Resources "
"Exploration - National Science and Technology Major Project (Program "
"No.2025ZD10069).")

# ---- References ----
h1(doc, "References")
refs = [
 "Warmerdam F (2008) The Geospatial Data Abstraction Library. In: Hall GB, Leahy MG (eds) Open Source Approaches in Spatial Data Handling. Advances in Geographic Information Science. Springer, Berlin, Heidelberg, pp 87–104. https://doi.org/10.1007/978-3-540-74831-1_5",
 "GDAL/OGR contributors (2026) Vector drivers list. https://gdal.org/drivers/vector/index.html. Accessed 22 July 2026.",
 "Li S (2022) pymapgis: a Python library for reading MapGIS 6.x/67 vector files. https://github.com/leecugb/pymapgis. Accessed 22 July 2026.",
 "MathsionYang (2019) MapGIS2ArcGIS. https://github.com/MathsionYang/MapGIS2ArcGIS. Accessed 22 July 2026.",
 "BenChao1998 (2025) ConvertMapGIS. https://github.com/BenChao1998/ConvertMapGIS. Accessed 22 July 2026.",
 "Wong W (2021) OGC.net. https://github.com/WenboWong/OGC.net. Accessed 22 July 2026.",
 "Wang X (2013) Data analysis and reading test on MapGIS point file. Science of Surveying and Mapping (in Chinese) 2013(1):112–115.",
 "Wu L, Feng J (2006) Study on the key techniques of MapGIS and Geodatabase data format conversion. Chinese Journal of Engineering Geophysics (in Chinese).",
 "Chen H, Wu J, Wang J (2000) Realisation of format conversion in MapGIS with MATLAB. Computing Techniques for Geophysical and Geochemical Exploration (in Chinese) 22(4):351–355.",
 "Shotwell M (2011) sas7bdat: sas7bdat reverse engineering documentation. CRAN (The R Foundation). https://doi.org/10.32614/cran.package.sas7bdat",
 "Senanayake AI, Chandler RJ, Daly T, Lewis E (2022) python-ags4: a Python library to read, write, and validate AGS4 geodata files. Journal of Open Source Software 7(79):4569. https://doi.org/10.21105/joss.04569",
 "Wright D, Wood R, Sylvander B (1998) ArcGMT: a suite of tools for conversion between Arc/INFO and Generic Mapping Tools (GMT). Computers & Geosciences 24(8):737–744. https://doi.org/10.1016/s0098-3004(98)00067-3",
 "Sen M, Duffy T (2005) GeoSciML: development of a generic GeoScience Markup Language. Computers & Geosciences 31(9):1095–1103. https://doi.org/10.1016/j.cageo.2004.12.003",
 "van den Bos J (2014) Lightweight runtime reverse engineering of binary file format variants. In: 2014 Software Evolution Week – IEEE CSMR-WCRE. IEEE, pp 367–370. https://doi.org/10.1109/csmr-wcre.2014.6747196",
 "Ma Y, Wang J, Xie S (2012) Analytical application of MapGIS for quality control in geological map spatial database constructing. Geo-information Science 13(6):758–762. https://doi.org/10.3724/sp.j.1047.2011.00758",
 "Han K, Pang J, Lu Y, Ding D, Fan B, Ju Y, Wang Z (2012) Research on sharing of geological map spatial data network under the \"OneGeology\" project: taking China 1:1M geological map data in MapGIS format as an example. Geo-information Science 13(6):742–749. https://doi.org/10.3724/sp.j.1047.2011.00742",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. ").bold = True
    p.add_run(r)

doc.save("paper.docx")
print("wrote paper.docx")
